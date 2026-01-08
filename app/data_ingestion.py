"""
Data Ingestion Module

This module handles reading and parsing admission documents (DOCX and TXT files).
It extracts text content from various formats and prepares it for processing.
"""

import os
from pathlib import Path
from typing import List, Optional
import docx
from app.config import DATA_DIR


def read_txt_file(file_path: Path) -> str:
    """
    Read text content from a TXT file.
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        Extracted text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def read_docx_file(file_path: Path) -> str:
    """
    Read text content from a DOCX file.
    
    Extracts text from all paragraphs in the document.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    doc = docx.Document(file_path)
    text_parts = []
    
    # Extract text from all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract text from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    
    return "\n".join(text_parts)


def load_admission_data(data_file: Optional[str] = None) -> str:
    """
    Load admission data from a file.
    
    Automatically detects file type (DOCX or TXT) and extracts text.
    If no file is specified, looks for the first DOCX or TXT file in the data directory.
    
    Args:
        data_file: Optional filename. If None, searches for files in data directory.
        
    Returns:
        Extracted text content
        
    Raises:
        FileNotFoundError: If no valid data file is found
        ValueError: If file format is not supported
    """
    if data_file:
        file_path = DATA_DIR / data_file
    else:
        # Find first available DOCX or TXT file
        docx_files = list(DATA_DIR.glob("*.docx"))
        txt_files = list(DATA_DIR.glob("*.txt"))
        
        if docx_files:
            file_path = docx_files[0]
        elif txt_files:
            file_path = txt_files[0]
        else:
            raise FileNotFoundError(
                f"No admission data file found in {DATA_DIR}. "
                "Please place a DOCX or TXT file in the data/ directory."
            )
    
    if not file_path.exists():
        raise FileNotFoundError(f"Data file not found: {file_path}")
    
    # Determine file type and read accordingly
    if file_path.suffix.lower() == '.docx':
        return read_docx_file(file_path)
    elif file_path.suffix.lower() == '.txt':
        return read_txt_file(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_path.suffix}. Use .docx or .txt")


def get_available_data_files() -> List[str]:
    """
    Get list of available data files in the data directory.
    
    Returns:
        List of filenames (DOCX and TXT files)
    """
    files = []
    for ext in ['*.docx', '*.txt']:
        files.extend([f.name for f in DATA_DIR.glob(ext)])
    return files






